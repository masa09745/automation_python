import tkinter as tk
import tkinter.ttk as ttk
import tkinter.messagebox as mb

from tkinter import messagebox

import glob
import os
import openpyxl as op
import docx as dx
import sys
import re
import math


def create():
  ship_input=ship_text.get()
  ech_input=ech_text.get()

  if getattr(sys, "frozen", False):
    os.chdir(os.path.dirname(os.path.abspath(sys.executable)))
    data_dir = os.path.join(os.path.dirname(sys.executable), "Data")
    dir = os.getcwd()
  else:
    data_dir = os.path.join(os.path.dirname(__file__), "Data")
    dir = os.getcwd()

  new_dir_filepath = dir + f'/JA{ship_input}_{ech_input}'
  os.mkdir(new_dir_filepath)

  select_file = ','.join(glob.glob(data_dir + '/*(ADV).xlsx' ))
  job_list = op.load_workbook(select_file)
  worksheets = job_list.sheetnames
  need_sheets = ['RTN', 'COA', 'EV']
  for worksheet in worksheets:
    flag = False
    for need_sheet in need_sheets:
      if worksheet == need_sheet:
        flag = False
        break
      else:
        flag = True
    if flag == True:
      sheet_count = len(job_list.sheetnames)
      if sheet_count >= 2:
        del job_list[worksheet]

  job_RTN = job_list['RTN']
  job_COA = job_list['COA']
  job_EV = job_list['EV']

  order_list = op.Workbook()
  count = 3

  for i in range(1, count+1):
    sheet = order_list.create_sheet()

  order_RTN = order_list["Sheet"]
  order_RTN.title = "RTN"
  order_REQ = order_list["Sheet1"]
  order_REQ.title = "REQ JOB"
  order_COA = order_list["Sheet2"]
  order_COA.title = "COA"
  order_EV = order_list["Sheet3"]
  order_EV.title = "EV"

  for i in range(3, job_RTN.max_row + 1):
    if job_RTN.cell(row = i, column =2).value == "REQUEST JOB CARD":
      order1 = [job_RTN.cell(row = i, column = 3).value]
      order_REQ.cell(row = 1, column = 1, value = "Order")
      order_REQ.append(order1)
    else:
      order2 = [job_RTN.cell(row = i, column = 3).value]
      order_RTN.append(order2)

  for i in range(3, job_COA.max_row + 1):
    order3 = [job_COA.cell(row = i, column = 3).value]
    order_COA.append(order3)

  for i in range(3, job_EV.max_row + 1):
    order4 = [job_EV.cell(row = i, column = 3).value]
    order_EV.append(order4)

  filename = f'{new_dir_filepath}/Order List.xlsx'
  order_list.save(filename)

  change_word = [ ['機番', ship_input], ['エチロン', ech_input],]

  change_icon1 = [ ['①', '■'], ['②', '□'], ['③', '□'], ['④', '□'], ['⑤', '□'], ['⑥', '□'],]
  change_icon2 = [ ['①', '□'], ['②', '□'], ['③', '□'], ['④', '□'], ['⑤', '■'], ['⑥', '□'],]
  change_icon3 = [ ['①', '□'], ['②', '□'], ['③', '■'], ['④', '□'], ['⑤', '□'], ['⑥', '□'],]
  change_icon4 = [ ['①', '□'], ['②', '□'], ['③', '□'], ['④', '■'], ['⑤', '□'], ['⑥', '□'],]

  os.makedirs(f'{new_dir_filepath}/REQ JOB')
  os.makedirs(f'{new_dir_filepath}/COA')
  os.makedirs(f'{new_dir_filepath}/EV')

  #作業アサインシート原紙の作成
  word_file = ','.join(glob.glob(data_dir + '/作業アサインシート原紙用.docx'))
  doc = dx.Document(word_file)

  tbl = doc.tables[0]
  target = tbl.rows[2]
  for cell in target.cells:
    cell_para = cell.paragraphs[0]
    for run in cell_para.runs:
      for i in range(len(change_word)):
        run.text = re.sub(change_word[i][0], change_word[i][1], run.text)
  rename_file = f'{new_dir_filepath}/JA{ship_input}_{ech_input}.docx'
  doc.save(rename_file)

  order_list = op.load_workbook(f'{new_dir_filepath}/Order List.xlsx')

  ws1 = order_list['RTN']
  value1 = [[cell.value for cell in row1] for row1 in ws1]
  ws2 = order_list['REQ JOB']
  value2 = [[cell.value for cell in row1] for row1 in ws2]
  ws3 = order_list['COA']
  value3 = [[cell.value for cell in row1] for row1 in ws3]
  ws4 = order_list['EV']
  value4 = [[cell.value for cell in row1] for row1 in ws4]

  word_file = ','.join(glob.glob(data_dir + '/作業アサインシート.docx'))

  dir_count = math.ceil((len(value1)-1)/100)

  for i in range(dir_count):
    i=i+1
    if i==dir_count:
      os.makedirs(f'{new_dir_filepath}/RTN_{len(value1)-1}まで')
    else:
      os.makedirs(f'{new_dir_filepath}/RTN_{i*100}まで')

  for y in range(1, len(value1)):
    dic1 = dict(Index = str(y))
    dic2 = dict(zip(value1[0], value1[y]))
    doc = dx.Document(word_file)

    for sec in doc.sections:
      for para in sec.header.paragraphs:
        for key,value in dic1.items():
          para.text = para.text.replace(key, str(value))

    tbl = doc.tables[0]
    target = tbl.rows[2]
    for cell in target.cells:
      cell_para = cell.paragraphs[0]
      for run in cell_para.runs:
        for i in range(len(change_word)):
          run.text = re.sub(change_word[i][0], change_word[i][1], run.text)

    target = tbl.rows[3]
    for cell in target.cells:
      cell_para = cell.paragraphs[0]
      for run in cell_para.runs:
        for x in range(len(change_icon1)):
          run.text = re.sub(change_icon1[x][0], change_icon1[x][1], run.text)
        for key, value in dic2.items():
          run.text = run.text.replace(key, str(value))

    if dir_count*100-100 < y <= dir_count*100:
      word_newFilePath = f'{new_dir_filepath}/RTN_{len(value1)-1}まで/{y}_{value1[y][0]}.docx'
      doc.save(word_newFilePath)
    elif y <= dir_count*100-100:
      for i in range(dir_count-1):
        if (i+1)*100-99 <= y <= (i+1)*100:
          word_newFilePath = f'{new_dir_filepath}/RTN_{(i+1)*100}まで/{y}_{value1[y][0]}.docx'
          doc.save(word_newFilePath)

  for y in range(1, len(value2)):
    dic1 = dict(Index = str(y))
    dic2 = dict(zip(value2[0], value2[y]))
    doc = dx.Document(word_file)

    for sec in doc.sections:
      for para in sec.header.paragraphs:
        for key,value in dic1.items():
          para.text = para.text.replace(key, str(value))

    tbl = doc.tables[0]
    target = tbl.rows[2]
    for cell in target.cells:
      cell_para = cell.paragraphs[0]
      for run in cell_para.runs:
        for i in range(len(change_word)):
          run.text = re.sub(change_word[i][0], change_word[i][1], run.text)

    target = tbl.rows[3]
    for cell in target.cells:
      cell_para = cell.paragraphs[0]
      for run in cell_para.runs:
        for x in range(len(change_icon2)):
          run.text = re.sub(change_icon2[x][0], change_icon2[x][1], run.text)
        for key, value in dic2.items():
          run.text = run.text.replace(key, str(value))
        word_newFilePath = f'{new_dir_filepath}/REQ JOB/{y}_{value2[y][0]}.docx'
        doc.save(word_newFilePath)

  for y in range(1, len(value3)):
    dic1 = dict(Index = str(y))
    dic2 = dict(zip(value3[0], value3[y]))
    doc = dx.Document(word_file)

    for sec in doc.sections:
      for para in sec.header.paragraphs:
        for key,value in dic1.items():
          para.text = para.text.replace(key, str(value))

    tbl = doc.tables[0]
    target = tbl.rows[2]
    for cell in target.cells:
      cell_para = cell.paragraphs[0]
      for run in cell_para.runs:
        for i in range(len(change_word)):
          run.text = re.sub(change_word[i][0], change_word[i][1], run.text)

    target = tbl.rows[3]
    for cell in target.cells:
      cell_para = cell.paragraphs[0]
      for run in cell_para.runs:
        for x in range(len(change_icon3)):
          run.text = re.sub(change_icon3[x][0], change_icon3[x][1], run.text)
        for key, value in dic2.items():
          run.text = run.text.replace(key, str(value))
        word_newFilePath = f'{new_dir_filepath}/COA/{y}_{value3[y][0]}.docx'
        doc.save(word_newFilePath)

  for y in range(1, len(value4)):
    dic1 = dict(Index = str(y))
    dic2 = dict(zip(value4[0], value4[y]))
    doc = dx.Document(word_file)

    for sec in doc.sections:
      for para in sec.header.paragraphs:
        for key,value in dic1.items():
          para.text = para.text.replace(key, str(value))

    tbl = doc.tables[0]
    target = tbl.rows[2]
    for cell in target.cells:
      cell_para = cell.paragraphs[0]
      for run in cell_para.runs:
        for i in range(len(change_word)):
          run.text = re.sub(change_word[i][0], change_word[i][1], run.text)

    target = tbl.rows[3]
    for cell in target.cells:
      cell_para = cell.paragraphs[0]
      for run in cell_para.runs:
        for x in range(len(change_icon4)):
          run.text = re.sub(change_icon4[x][0], change_icon4[x][1], run.text)
        for key, value in dic2.items():
          run.text = run.text.replace(key, str(value))
        word_newFilePath = f'{new_dir_filepath}/EV/{y}_{value4[y][0]}.docx'
        doc.save(word_newFilePath)




def click():
  create()
  messagebox.showinfo("完了", "完了しました")
  
def close():
  main_win.destroy()



# GUIのレイアウトコード
main_win = tk.Tk()
main_win.geometry('300x200')
main_win.title('作業アサインシート一括印刷')

main_frame = ttk.Frame(main_win)
main_frame.grid(column=0, row=0, sticky=tk.NSEW, padx=20, pady=20)

ship_label = tk.Label(main_frame, text='機番 (JAは不要)')
ship_text = tk.StringVar()
shiptext_entry = tk.Entry(main_frame, textvariable=ship_text)

ech_label = tk.Label(main_frame, text='エチロン')
ech_text = tk.StringVar()
echtext_entry = tk.Entry(main_frame, textvariable=ech_text)

create_btn = tk.Button(main_frame, text="作成", command=click)
close_btn = tk.Button(main_frame, text="閉じる", command=close)

ship_label.grid(column=0, row=0, sticky=tk.W)
shiptext_entry.grid(column=0, row=1, pady=5, sticky=tk.EW)

ech_label.grid(column=0, row=2, sticky=tk.W)
echtext_entry.grid(column=0, row=3, pady=5, sticky=tk.EW)

create_btn.grid(column=0, row=4, sticky=tk.W, pady=10)
close_btn.grid(column=0, row=4, sticky=tk.E, pady=10)

main_win.columnconfigure(0, weight=1)
main_win.rowconfigure(0, weight=1)
main_frame.columnconfigure(0, weight=1)

main_win.mainloop()
