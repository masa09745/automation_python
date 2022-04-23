import tkinter as tk
import tkinter.ttk as ttk
import tkinter.messagebox as mb

from tkinter import messagebox

import glob
import os
import openpyxl as op
import docx as dx
import sys
import re
import math


def create():
  ship_input=ship_text.get()
  ech_input=ech_text.get()
  
  dir = os.getcwd()
  new_dir_filepath = dir + f'/JA{ship_input}_{ech_input}'
  os.makedirs(new_dir_filepath)
  select_file = ','.join(glob.glob(dir + '/Data/*(ADV)*.xlsx'))
  
  joblist = op.load_workbook(select_file)
  worksheets = joblist.sheetnames
  needsheets = ['RTN', 'COA', 'EV']

  for worksheet in worksheets:
    flag = False
    for needsheet in needsheets:
      if worksheet == needsheet:
        flag = False
        break
      else:
        flag = True
    if flag == True:
      sheetcount = len(joblist.sheetnames)
      if sheetcount >= 2:
        del joblist[worksheet]
  filename = f'{new_dir_filepath}/Job List.xlsx'
  joblist.save(filename)
  
  new_joblist = op.load_workbook(f'{new_dir_filepath}/Job List.xlsx')
  job_RTN = new_joblist['RTN']
  job_COA = new_joblist['COA']
  job_EV = new_joblist['EV']
  
  orderlist = op.Workbook()
  count = 3

  for i in range(1,count+1):
    sheet = orderlist.create_sheet()
    
  order_RTN = orderlist["Sheet"]
  order_RTN.title = "RTN"
  order_REQ = orderlist["Sheet1"]
  order_REQ.title = "REQ JOB"
  order_COA = orderlist["Sheet2"]
  order_COA.title = "COA"
  order_EV = orderlist["Sheet3"]
  order_EV.title = "EV"
  
  for i in range(3, job_RTN.max_row + 1):
    if job_RTN.cell(row = i, column = 2).value == 'REQUEST JOB CARD':
      order1 = [job_RTN.cell(row = i, column = 3).value]
      order_REQ.cell(row = 1, column = 1, value = "Order")
      order_REQ.append(order1)
    else:
      order2 = [job_RTN.cell(row = i, column = 3).value]
      order_RTN.append(order2)

  for i in range(3, job_COA.max_row + 1):
    order3 = [job_COA.cell(row = i, column = 3).value]
    order_COA.append(order3)

  for i in range(3, job_EV.max_row + 1):
    order4 = [job_EV.cell(row = i, column = 3).value]
    order_EV.append(order4)
  
  filename = f'{new_dir_filepath}/Order List.xlsx'
  orderlist.save(filename)
  
  change_word = [
    ['機番', ship_input],
    ['エチロン', ech_input],
  ]

  

  orderlist = op.load_workbook(f'{new_dir_filepath}/Order List.xlsx')
  ws1 = orderlist['RTN']
  
  value1 = [[cell.value for cell in row1] for row1 in ws1]
  change_icon1 = [ ['①', '■'], ['②', '□'], ['③', '□'], ['④', '□'], ['⑤', '□'], ['⑥', '□'],]
  
  word_file = ','.join(glob.glob(dir + '/data/作業アサインシート.docx'))

  sum =len(value1)/100
  test = math.ceil(sum)
  print(test)
  for a in range(test):
    b = a+1
    os.makedirs(f'{new_dir_filepath}/RTN_{b}')
  
  os.makedirs(f'{new_dir_filepath}/RTN(1-100)')
  
  for y in range(1, len(value1)):
    if y <= 100:
      
      dic1 = dict(Index = str(y))
      dic2 = dict(zip(value1[0], value1[y]))
      doc = dx.Document(word_file)
      
        
      for sec in doc.sections:
        for para in sec.header.paragraphs:
          for key,value in dic1.items():
            para.text = para.text.replace(key, str(value))

      tbl = doc.tables[0]
      target = tbl.rows[2]
      for cell in target.cells:
        cell_para = cell.paragraphs[0]
        for run in cell_para.runs:
          for i in range(len(change_word)):
            run.text = re.sub(change_word[i][0], change_word[i][1], run.text)

      target = tbl.rows[3]
      for cell in target.cells:
        cell_para = cell.paragraphs[0]
        for run in cell_para.runs:
          for x in range(len(change_icon1)):
            run.text = re.sub(change_icon1[x][0], change_icon1[x][1], run.text)
          for key, value in dic2.items():
            run.text = run.text.replace(key, str(value))
      word_newFilePath = f'{new_dir_filepath}/RTN(1-100)/{y}_{value1[y][0]}.docx'
      doc.save(word_newFilePath)
  
def click():
  create()
  messagebox.showinfo("完了", "完了しました")

def close():
  main_win.destroy()
  

main_win = tk.Tk()
main_win.geometry('300x200')
main_win.title('作業アサインシート一括印刷')

main_frame = ttk.Frame(main_win)
main_frame.grid(column=0, row=0, sticky=tk.NSEW, padx=20, pady=20)

ship_label = tk.Label(main_frame, text='機番 (JAは不要)')
ship_text = tk.StringVar()
shiptext_entry = tk.Entry(main_frame, textvariable=ship_text)

ech_label = tk.Label(main_frame, text='エチロン')
ech_text = tk.StringVar()
echtext_entry = tk.Entry(main_frame, textvariable=ech_text)

create_btn = tk.Button(main_frame, text="作成", command=click)
close_btn = tk.Button(main_frame, text="閉じる", command=close)

ship_label.grid(column=0, row=0, sticky=tk.W)
shiptext_entry.grid(column=0, row=1, pady=5, sticky=tk.EW)

ech_label.grid(column=0, row=2, sticky=tk.W)
echtext_entry.grid(column=0, row=3, pady=5, sticky=tk.EW)

create_btn.grid(column=0, row=4, sticky=tk.W, pady=10)
close_btn.grid(column=0, row=4, sticky=tk.E, pady=10)

main_win.columnconfigure(0, weight=1)
main_win.rowconfigure(0, weight=1)
main_frame.columnconfigure(0, weight=1)

main_win.mainloop()