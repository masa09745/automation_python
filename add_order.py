import openpyxl as op
import docx as dx
import os
import glob
import re


orderlist = op.load_workbook('Order List.xlsx')
ws1 = orderlist['RTN']
value1 = [[cell.value for cell in row1] for row1 in ws1]

new_dir_filepath = 'test'
os.makedirs(new_dir_filepath)

for i in range(1, len(value1)):
  dic = dict(zip(value1[0], value1[i]))

  word_file = glob.glob('*.docx')
  select_file = ','.join(word_file)
  doc = dx.Document(select_file)
  tbl = doc.tables[0]
  target = tbl.rows[3]
  for cell in target.cells:
    cell_para = cell.paragraphs[0]
    for run in cell_para.runs:
      for key, value in dic.items():
        run.text = run.text.replace(key, str(value))

      word_newFilePath = f'test/{i}_{value1[i][0]}.docx'
      doc.save(word_newFilePath)
