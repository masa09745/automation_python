import tkinter as tk
import tkinter.ttk as ttk
import tkinter.messagebox as mb

from tkinter import messagebox

import openpyxl as op
import docx as dx
import glob
import re
import os

# 入力した機番、エチロンからフォルダの作成
def makedirs():
  ship_input=ship_text.get()
  ech_input=ech_text.get()
  
  new_dir_filepath = f'JA{ship_input}_{ech_input}'
  os.makedirs(new_dir_filepath)
  
# ADV版の計画表から不要なシートの削除
def del_sheet():
  ship_input=ship_text.get()
  ech_input=ech_text.get()

  excel_files = glob.glob('必要データ/*.xlsx')
  select_file = ','.join(excel_files)
  joblist = op.load_workbook(select_file)
  worksheets = joblist.sheetnames

  needsheets = ['RTN', 'COA', 'EV', 'IO']

  for worksheet in worksheets:
    flag = False
    for needsheet in needsheets:
      if worksheet == needsheet:
        flag = False
        break
      else:
        flag = True
    if flag == True:
      sheetcount = len(joblist.sheetnames)
      if sheetcount >= 2:
        del joblist[worksheet]
  filename = f'JA{ship_input}_{ech_input}/Job List.xlsx'
  joblist.save(filename)
  
# Order ListのExcelファイル作成
def mkfile():
  ship_input=ship_text.get()
  ech_input=ech_text.get()
  
  orderlist = op.Workbook()
  count = 3

  for i in range(1,count+1):
    sheet = orderlist.create_sheet()

  RTN_Order = orderlist["Sheet"]
  RTN_Order.title = "RTN"
  COA_Order = orderlist["Sheet1"]
  COA_Order.title = "COA"
  EV_Order = orderlist["Sheet2"]
  EV_Order.title = "EV"
  IO_Order = orderlist["Sheet3"]
  IO_Order.title = "IO"
  
  filename = f'JA{ship_input}_{ech_input}/Order List.xlsx'
  orderlist.save(filename)
  
# Job ListからOrder Numberを貼り付けて Order Listの作成
def copy_order():
  ship_input=ship_text.get()
  ech_input=ech_text.get()  
  
  newlist = op.load_workbook(f'JA{ship_input}_{ech_input}/Job List.xlsx')
  new_RTN = newlist['RTN']
  new_COA = newlist['COA']
  new_EV = newlist['EV']
  new_IO = newlist['IO']

  orderlist = op.load_workbook(f'JA{ship_input}_{ech_input}/Order List.xlsx')
  order_RTN = orderlist['RTN']
  order_COA = orderlist['COA']
  order_EV = orderlist['EV']
  order_IO = orderlist['IO']
  
  for i in range(3, new_RTN.max_row + 1):
    order1 = new_RTN.cell(row = i, column = 3).value
    order_RTN.cell(row = i-2, column = 1, value = order1)
    order2 = new_COA.cell(row = i, column = 3).value
    order_COA.cell(row = i-2, column = 1, value = order2)
    order3 = new_EV.cell(row = i, column = 3).value
    order_EV.cell(row = i-2, column = 1, value = order3)
    order4 = new_IO.cell(row = i, column = 3).value
    order_IO.cell(row = i-2, column = 1, value = order4)
    
  filename = f'JA{ship_input}_{ech_input}/Order List.xlsx'
  orderlist.save(filename)

# 作業アサインシートの原紙の作成
def create():
  ship_input=ship_text.get()
  ech_input=ech_text.get()
  
  change_word = [
    ['機番', ship_input],
    ['エチロン', ech_input],
  ]
  
  word_file = glob.glob('必要データ/作業アサインシート原紙用.docx')
  select_file = ','.join(word_file)
  doc = dx.Document(select_file)

  tbl = doc.tables[0]
  target = tbl.rows[2]
  for cell in target.cells:
    cell_para = cell.paragraphs[0]
    for run in cell_para.runs:
      for i in range(len(change_word)):
        run.text = re.sub(change_word[i][0], change_word[i][1], run.text)
  rename_file = f'JA{ship_input}_{ech_input}/JA{ship_input}_{ech_input}.docx'
  doc.save(rename_file)

# 作業アサインシートへのOrder Numberの貼り付け
def paste():
  ship_input=ship_text.get()
  ech_input=ech_text.get()

  new_dir_filepath = 'RTN'
  os.makedirs(f'JA{ship_input}_{ech_input}/{new_dir_filepath}')

  orderlist = op.load_workbook('Order List.xlsx')
  ws1 = orderlist['RTN']
  value1 = [[cell.value for cell in row1] for row1 in ws1]

  for i in range(1, len(value1)):
    dic = dict(zip(value1[0], value1[i]))

    word_file = glob.glob('必要データ/作業アサインシート.docx')
    select_file = ','.join(word_file)
    doc = dx.Document(select_file)
    tbl = doc.tables[0]
    target = tbl.rows[3]
    for cell in target.cells:
      cell_para = cell.paragraphs[0]
      for run in cell_para.runs:
        for key, value in dic.items():
          run.text = run.text.replace(key, str(value))

        word_newFilePath = f'JA{ship_input}_{ech_input}/{new_dir_filepath}/{i}_{value1[i][0]}.docx'
        doc.save(word_newFilePath)

# 作成ボタンクリック時の挙動
def click():
  makedirs()
  del_sheet()
  mkfile()
  copy_order()
  create()
  paste()
  messagebox.showinfo("完了", "完了しました")


# GUIのレイアウトコード
main_win = tk.Tk()
main_win.geometry('300x200')
main_win.title('作業アサインシート一括印刷')

main_frame = ttk.Frame(main_win)
main_frame.grid(column=0, row=0, sticky=tk.NSEW, padx=20, pady=20)

ship_label = tk.Label(main_frame, text='機番 (JAは不要)')
ship_text = tk.StringVar()
shiptext_entry = tk.Entry(main_frame, textvariable=ship_text)

ech_label = tk.Label(main_frame, text='エチロン')
ech_text = tk.StringVar()
echtext_entry = tk.Entry(main_frame, textvariable=ech_text)

create_btn = tk.Button(main_frame, text="作成", command=click)

ship_label.grid(column=0, row=0, sticky=tk.W)
shiptext_entry.grid(column=0, row=1, pady=5, sticky=tk.EW)

ech_label.grid(column=0, row=3, sticky=tk.W)
echtext_entry.grid(column=0, row=4, pady=5, sticky=tk.EW)

create_btn.grid(column=0, row=5, pady=15)

main_win.columnconfigure(0, weight=1)
main_win.rowconfigure(0, weight=1)
main_frame.columnconfigure(0, weight=1)

main_win.mainloop()