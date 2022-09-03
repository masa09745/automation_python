import tkinter as tk
import tkinter.ttk as ttk
import tkinter.messagebox as mb

from tkinter import messagebox

import openpyxl as op
import docx as dx
import os
import sys
import glob
import re

def create_rev():
  ship_input = ship_text.get()
  ech_input = ech_text.get()

  if getattr(sys, "frozen", False):
    os.chdir(os.path.dirname(os.path.dirname(os.path.abspath(sys.executable))))
    data_dir = os.path.join(os.path.dirname(sys.executable), "Data")
    dir = os.getcwd()
  else:
    data_dir = os.path.join(os.path.dirname(__file__), "Data")
    dir = os.getcwd()

  new_dir_filepath = dir + f'/JA{ship_input}_{ech_input}'

  if os.path.exists(new_dir_filepath):
    if not os.path.exists(f'{new_dir_filepath}/Rev Notice'):
      os.mkdir(f'{new_dir_filepath}/Rev Notice')

  else:
    os.mkdir(new_dir_filepath)
    os.mkdir(f'{new_dir_filepath}/Rev Notice')


  select_file = ','.join(glob.glob(data_dir + '/*Rev*.xlsx' ))
  wb = op.load_workbook(select_file)
  worksheets = wb.sheetnames
  ws = wb['RN']
  val = ws['A1'].value

  for i in range(ws.max_row, 3, -1):
    if ws.cell(row = i, column = 3).value == 'XLD':
      ws.delete_rows(i)
    if ws.cell(row = i, column = 1).value == None:
      ws.delete_rows(i)
  filename = f'{new_dir_filepath}/{val}.xlsx'
  wb.save(filename)

  rev_list = op.load_workbook(f'{new_dir_filepath}/{val}.xlsx')
  rev_data = rev_list['RN']

  word_file = ','.join(glob.glob(data_dir + '/*Rev*.docx'))
  doc = dx.Document(word_file)

  for i in range(4, rev_data.max_row + 1):
    doc = dx.Document(word_file)
    change_word = [ ['機番', ship_input], ['エチロン', ech_input],]
    change_icon1 = [ ['①', '■'], ['②', '□'], ['③', '□'], ['④', '□'], ['⑤', '□'], ['⑥', '□'],]
    change_icon2 = [ ['①', '□'], ['②', '□'], ['③', '■'], ['④', '□'], ['⑤', '□'], ['⑥', '□'],]
    change_icon3 = [ ['①', '□'], ['②', '□'], ['③', '□'], ['④', '■'], ['⑤', '□'], ['⑥', '□'],]

    value1 = rev_data.cell(row = i, column = 1).value
    dic1 = dict(Index = value1)

    value2 = rev_data.cell(row = i, column = 5).value
    dic2 = dict(Order = value2)

    value3 = rev_data.cell(row = i, column = 4 ).value

    for sec in doc.sections:
      for para in sec.header.paragraphs:
        for key, value in dic1.items():
          para.text = re.sub(key, str(value), para.text)

    tbl = doc.tables[0]
    for cell in tbl.rows[2].cells:
      cell_para = cell.paragraphs[0]
      for run in cell_para.runs:
        for i in range(len(change_word)):
          run.text = re.sub(change_word[i][0], change_word[i][1], run.text)

    for cell in tbl.rows[3].cells:
      cell_para = cell.paragraphs[0]
      for run in cell_para.runs:
        if re.search('COA', value3):
          for i in range(len(change_icon2)):
            run.text = re.sub(change_icon2[i][0], change_icon2[i][1], run.text)
        elif re.search('EV', value3):
          for i in range(len(change_icon3)):
            run.text = re.sub(change_icon3[i][0], change_icon3[i][1], run.text)
        else:
          for i in range(len(change_icon1)):
            run.text = re.sub(change_icon1[i][0], change_icon1[i][1], run.text)
        for key, value in dic2.items():
          run.text = re.sub(key, str(value), run.text)
    word_newFilePath = f'{new_dir_filepath}/Rev Notice/{value1}_{value2}.docx'
    doc.save(word_newFilePath)
    
def click():
  create_rev()
  messagebox.showinfo("完了", "完了しました")

def close():
  main_win.destroy()
  

main_win = tk.Tk()
main_win.geometry('300x200')
main_win.title('作業アサインシート一括印刷')

main_frame = ttk.Frame(main_win)
main_frame.grid(column=0, row=0, sticky=tk.NSEW, padx=20, pady=20)

ship_label = tk.Label(main_frame, text='機番 (JAは不要)')
ship_text = tk.StringVar()
shiptext_entry = tk.Entry(main_frame, textvariable=ship_text)

ech_label = tk.Label(main_frame, text='エチロン')
ech_text = tk.StringVar()
echtext_entry = tk.Entry(main_frame, textvariable=ech_text)

create_btn = tk.Button(main_frame, text="作成", command=click)
close_btn = tk.Button(main_frame, text="閉じる", command=close)

ship_label.grid(column=0, row=0, sticky=tk.W)
shiptext_entry.grid(column=0, row=1, pady=5, sticky=tk.EW)

ech_label.grid(column=0, row=2, sticky=tk.W)
echtext_entry.grid(column=0, row=3, pady=5, sticky=tk.EW)

create_btn.grid(column=0, row=4, sticky=tk.W, pady=10)
close_btn.grid(column=0, row=4, sticky=tk.E, pady=10)

main_win.columnconfigure(0, weight=1)
main_win.rowconfigure(0, weight=1)
main_frame.columnconfigure(0, weight=1)

main_win.mainloop()